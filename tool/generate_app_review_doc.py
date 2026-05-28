from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(r"D:\APP\flutter_application_1")
OUTPUT = ROOT / f"伴一下APP项目整理报告_{date.today().isoformat()}.docx"


def count_files(path: Path, pattern: str) -> int:
    return sum(1 for _ in path.rglob(pattern))


def count_text(path: Path, pattern: str) -> int:
    return path.read_text(encoding="utf-8").count(pattern)


def add_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False, color: str | None = None) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold = bold
    font = run.font
    font.size = Pt(10.5)
    font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    if color:
        font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def add_bullet(document: Document, text: str) -> None:
    p = document.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(10.5)


def add_number(document: Document, text: str) -> None:
    p = document.add_paragraph(style="List Number")
    run = p.add_run(text)
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(10.5)


def set_run_font(run, size: float = 10.5, bold: bool = False, color: str | None = None) -> None:
    font = run.font
    font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    font.size = Pt(size)
    font.bold = bold
    if color:
        font.color.rgb = RGBColor.from_string(color)


def add_para(document: Document, text: str, style: str = "Normal", first_line_cm: float = 0.74) -> None:
    p = document.add_paragraph(style=style)
    p.paragraph_format.first_line_indent = Cm(first_line_cm)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.35
    run = p.add_run(text)
    set_run_font(run)


def add_reference_line(document: Document, refs: list[str]) -> None:
    p = document.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    label = p.add_run("代码依据：")
    set_run_font(label, bold=True, color="3D5AFE")
    body = p.add_run("；".join(refs))
    set_run_font(body, size=10)


def build_styles(document: Document) -> None:
    styles = document.styles

    normal = styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)

    for style_name, size, color in [
        ("Title", 22, "1F2A44"),
        ("Subtitle", 12, "5B657A"),
        ("Heading 1", 16, "2447B2"),
        ("Heading 2", 13, "2447B2"),
        ("Heading 3", 11.5, "2447B2"),
    ]:
        style = styles[style_name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = True if "Heading" in style_name or style_name == "Title" else False
        style.font.color.rgb = RGBColor.from_string(color)

    if "Callout" not in styles:
        callout = styles.add_style("Callout", WD_STYLE_TYPE.PARAGRAPH)
    else:
        callout = styles["Callout"]
    callout.font.name = "Microsoft YaHei"
    callout._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    callout.font.size = Pt(10.5)


def add_cover(document: Document) -> None:
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(90)
    run = p.add_run("伴一下 App 项目整理报告")
    set_run_font(run, size=24, bold=True, color="1F2A44")

    p2 = document.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p2.add_run("基于当前仓库代码、SQL、Edge Function 与项目文档的模块化盘点")
    set_run_font(run, size=12, color="5B657A")

    table = document.add_table(rows=4, cols=2)
    table.style = "Table Grid"
    table.autofit = True
    rows = [
        ("扫描日期", str(date.today())),
        ("项目路径", str(ROOT)),
        ("盘点口径", "以仓库内已存在代码与文档为准，不把未落地的设想算作已完成"),
        ("文档目的", "梳理已完成工作、识别剩余建设项，并给出面向正式运营的优先级路线图"),
    ]
    for row, (left, right) in zip(table.rows, rows):
        set_cell_text(row.cells[0], left, bold=True, color="2447B2")
        set_cell_text(row.cells[1], right)
        add_cell_shading(row.cells[0], "EAF0FF")

    document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def add_overview_metrics(document: Document, metrics: dict[str, str]) -> None:
    document.add_heading("一、项目总览", level=1)
    add_para(
        document,
        "本项目是一套以“广场内容 + 地陪服务 + 需求定制 + 订单支付 + 消息沟通 + 地陪审核”为主线的 Flutter 客户端，当前后端主要依赖 Supabase，地图接入基于高德 Web Service，支付链路已预埋支付宝 App 支付和 Supabase Edge Function，下游也保留了迁移到自建服务器的结构入口。",
    )
    add_para(
        document,
        "从仓库结构看，它已经不再是单纯的原型演示，而是处在“核心流程已具备雏形、若干模块可演示、但距离正式运营仍缺关键闭环”的阶段。最值得肯定的是，页面、Provider、SQL、Edge Function 和客户交付文档已经同步出现，说明项目已经开始从做页面转向做系统。",
    )

    table = document.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    set_cell_text(hdr[0], "扫描项", bold=True, color="FFFFFF")
    set_cell_text(hdr[1], "结果", bold=True, color="FFFFFF")
    add_cell_shading(hdr[0], "3D6CF5")
    add_cell_shading(hdr[1], "3D6CF5")
    for key, value in metrics.items():
        row = table.add_row().cells
        set_cell_text(row[0], key, bold=True, color="2447B2")
        set_cell_text(row[1], value)
        add_cell_shading(row[0], "F4F7FF")


def add_status_table(document: Document, rows_data: list[tuple[str, str, str, str]]) -> None:
    document.add_heading("二、核心模块状态总表", level=1)
    add_para(
        document,
        "下表用于先给项目定调。这里的“状态”不是产品想象，而是按当前代码落地程度来标记：`已具备演示闭环` 表示前端、Provider 或数据库已有真实支撑；`部分完成` 表示流程能走一段，但还缺关键规则、接口或运营能力；`待建设` 表示目前只有占位、文档或局部接口预留。",
    )

    table = document.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    headers = ["模块", "当前状态", "当前能做什么", "最关键缺口"]
    for cell, title in zip(table.rows[0].cells, headers):
        set_cell_text(cell, title, bold=True, color="FFFFFF")
        add_cell_shading(cell, "3D6CF5")

    for row_data in rows_data:
        row = table.add_row().cells
        for idx, value in enumerate(row_data):
            set_cell_text(row[idx], value)
            if idx == 0:
                add_cell_shading(row[idx], "F4F7FF")


def add_module_section(document: Document, module: dict[str, object]) -> None:
    document.add_heading(module["title"], level=1)
    add_para(document, module["summary"])

    document.add_heading("已完成工作", level=2)
    for item in module["done"]:
        add_bullet(document, item)

    document.add_heading("当前仍需继续的工作", level=2)
    for item in module["todo"]:
        add_bullet(document, item)

    document.add_heading("对正式运营的判断", level=2)
    add_para(document, module["judgement"])
    add_reference_line(document, module["refs"])


def add_process_section(document: Document, title: str, intro: str, items: list[str]) -> None:
    document.add_heading(title, level=1)
    add_para(document, intro)
    for item in items:
        add_number(document, item)


def add_priority_table(document: Document, items: list[tuple[str, str, str, str]]) -> None:
    document.add_heading("二十、建议执行优先级", level=1)
    add_para(
        document,
        "下面的优先级不是按“好不好看”排序，而是按“对真实运营是否卡脖子”排序。一个地陪类 App 真要上线，先卡住的永远是身份、交易、履约、申诉和后台，而不是再多做两个展示页。",
    )
    table = document.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    for cell, text in zip(table.rows[0].cells, ["优先级", "目标", "本阶段应完成内容", "依赖外部资源"]):
        set_cell_text(cell, text, bold=True, color="FFFFFF")
        add_cell_shading(cell, "3D6CF5")

    for item in items:
        row = table.add_row().cells
        for idx, value in enumerate(item):
            set_cell_text(row[idx], value)
            if idx == 0:
                add_cell_shading(row[idx], "F4F7FF")


def add_client_section(document: Document, items: list[str]) -> None:
    document.add_heading("二十一、客户侧必须提供的内容", level=1)
    add_para(
        document,
        "从当前仓库和已有交付文档看，开发侧已经尽量把代码先铺出来了，真正阻塞正式化推进的，更多是客户主体、第三方账号、服务器、合规材料和正式密钥。下面这份清单可以直接拿去和客户对接。",
    )
    for item in items:
        add_bullet(document, item)


def add_conclusion(document: Document) -> None:
    document.add_heading("二十二、结论", level=1)
    add_para(
        document,
        "综合判断，这个项目已经具备“继续往正式产品推进”的基础：主导航、核心业务页、Provider 层、SQL 脚本、Edge Function 和客户交付文档都已经成型，说明它不是一套只有界面没有骨架的半成品。",
    )
    add_para(
        document,
        "但它当前仍然更接近“高完成度测试版”，还不能直接视为可运营正式版。离上线真正差的，不是再补几个页面，而是把角色体系、需求撮合、订单状态机、支付回调、退款售后、实名审核、短信正式通道、客服申诉、财务后台和运维环境全部闭环。",
    )
    add_para(
        document,
        "如果只选一个下一阶段主目标，最合理的方向是：先把“需求/下单/支付/履约/售后”这条交易主链做实，再同步补“地陪审核/后台/风控/短信/服务器迁移”。这条链一旦闭合，项目就会从“能演示的 App”真正进入“能试运营的 App”。",
    )


def main() -> None:
    metrics = {
        "Flutter 页面文件数量": f"{count_files(ROOT / 'lib' / 'pages', '*.dart')} 个",
        "Provider 数量": f"{count_files(ROOT / 'lib' / 'providers', '*.dart')} 个",
        "Service 数量": f"{count_files(ROOT / 'lib' / 'services', '*.dart')} 个",
        "已注册路由数量": f"{count_text(ROOT / 'lib' / 'config' / 'app_router.dart', 'GoRoute(')} 个",
        "Supabase Edge Function 数量": f"{sum(1 for p in (ROOT / 'supabase' / 'functions').iterdir() if p.is_dir())} 个",
        "SQL 初始化/补丁脚本": f"{count_files(ROOT, '*.sql')} 个",
        "SQL 中出现的 CREATE TABLE 语句": "24 处（包含不同初始化/补丁脚本中的重复定义）",
        "SQL 中出现的 CREATE OR REPLACE FUNCTION 语句": "8 处",
        "当前项目主技术栈": "Flutter + Provider + GoRouter + Supabase + 高德 Web Service + 支付宝 App 支付",
        "当前定位": "高完成度测试版 / 具备继续产品化的基础",
    }

    status_rows = [
        ("基础架构与配置", "已具备演示闭环", "配置统一、启动装配、Provider/Router 基本成型", "缺正式多环境、密钥治理与持续集成"),
        ("登录与用户体系", "部分完成", "短信登录抽象、用户资料同步、管理员/地陪/封禁判断已接入", "正式短信通道、实名/KYC、黑名单和设备风控未完成"),
        ("广场内容与招募", "已具备演示闭环", "帖子列表、详情、点赞、收藏、评论、关注流、分享/招募发布已落地", "内容审核、举报、话题运营和存储治理不足"),
        ("服务页与地陪展示", "部分完成", "城市/搜索/分类、地陪列表、详情、点赞收藏足迹可用", "接单能力、排班、报价、服务半径和履约规则不足"),
        ("需求定制", "部分完成", "需求创建、列表、地点搜索、时间选择、Supabase 发布已打通", "需求详情、报名/抢单/报价/确认链路未闭环"),
        ("订单支付结算", "部分完成", "下单、订单列表、支付宝下单参数、支付结果回写、钱包/流水结构已铺", "支付回调验签、退款、售后、分账对账、超时关闭未闭环"),
        ("地陪申请与审核", "部分完成", "申请表单、多步骤审核材料、后台审核入口、权限拦截已有", "实名核验 API、材料存储、审核记录、角色体系和操作留痕不足"),
        ("消息与通知", "部分完成", "会话列表、聊天室、实时消息、未读数已落地", "系统通知、活动通知、客服消息与订单会话绑定未完成"),
        ("个人中心", "部分完成", "资料编辑、身份展示、订单入口、收藏足迹、设置页已形成", "更完整的资料字段、提现、发票、评价、成就体系未完善"),
        ("地图与定位", "部分完成", "地点搜索、地理编码、逆地理、当前位置、静态地图、A-Z 城市表已有", "Android 定位稳定性、真地图选点与 SDK 方案仍需继续联调"),
        ("数据库与服务端", "部分完成", "核心表、RLS、RPC、支付/短信 Edge Function 已形成底座", "结构统一、触发器梳理、生产服务器迁移和接口治理未完成"),
        ("运维与交付", "部分完成", "客户交付、部署与 API 接入文档已写出", "正式服务器、域名、备案、证书、CI/CD 和监控未落地"),
    ]

    modules = [
        {
            "title": "三、基础架构与配置",
            "summary": "当前工程已经从“单文件堆页面”的阶段走出来，形成了 `config / bootstrap / app / providers / services / pages / widgets` 这一套比较清晰的目录分层。App 根部位于 `lib/app/app.dart`，启动初始化位于 `lib/bootstrap/app_bootstrap.dart`，全局路由在 `lib/config/app_router.dart`，应用级配置集中在 `lib/config/app_config.dart`、`payment_config.dart`、`amap_config.dart`、`auth_config.dart`。这一层的价值不是好看，而是让后续接正式服务器、多环境和多支付方式时不至于从页面里到处抠代码。",
            "done": [
                "已经形成统一的 Provider 装配入口，`UserProvider`、`PostProvider`、`GuideProvider`、`DemandProvider`、`OrderProvider`、`MessageProvider`、`ApplicationProvider` 都在 App 根部装配。",
                "已经形成统一路由中心，当前仓库内登记了 23 个路由，覆盖登录、帖子、地陪、需求、订单、消息、设置、审核等主流程。",
                "地图、支付、鉴权配置已经收口到 `lib/config`，避免继续在多个页面和服务文件里散落常量。",
                "支付服务与短信服务都做了抽象，说明项目已经考虑过后续从 Supabase 过渡到自建后端的需求。",
                "项目根目录已累计形成多份辅助文档，包含架构说明、Gap Review、客户交付步骤、部署与 API 接入说明。",
            ],
            "todo": [
                "需要补正式环境与测试环境的区分机制，例如 `dev / staging / prod` 三套配置，不应继续把默认值直接写死在 `AppConfig`。",
                "需要把敏感值从仓库默认值中彻底抽离，尤其是 Supabase URL、Anon Key、高德 Key 等，不宜长期以默认常量形式存在。",
                "需要补持续集成和构建检查，例如统一的 `analyze / test / build` 流程，否则后续多人协作容易把结构再次拉散。",
                "需要整理一份真正面向开发者的 README，目前根目录 README 仍是 Flutter 默认模板，无法作为项目文档入口。",
            ],
            "judgement": "这一层已经具备继续演进的基础，说明工程骨架不是问题；真正的问题在于它还没有进入正式项目常见的环境治理、密钥治理和自动化治理阶段。",
            "refs": [
                "lib/app/app.dart",
                "lib/config/app_router.dart",
                "lib/config/app_config.dart",
                "ARCHITECTURE_OVERVIEW.md",
            ],
        },
        {
            "title": "四、登录鉴权与用户体系",
            "summary": "登录模块当前以手机号短信登录为主，Flutter 侧通过 `PhoneAuthService` 做了抽象，默认实现是 Supabase Phone Auth。用户信息不只停留在 Auth 用户对象，而是会同步到 `public.users` 表，再由 `UserProvider` 汇总用户资料、管理员状态、地陪状态、封禁状态、关注/粉丝数量和申请状态。这说明项目已经开始从“账号”过渡到“用户实体”。",
            "done": [
                "已实现短信验证码发送和校验的抽象接口，未来可以从 Supabase 平滑替换到阿里云等国内短信服务。",
                "已实现登录后用户与 `users` 表同步，用户首次登录会自动补建基础资料。",
                "已在 `UserProvider` 中聚合管理员、地陪、封禁、申请中等角色上下文，不再让页面各自猜身份。",
                "已扩展用户资料字段，包括昵称、头像、简介、性别、城市、生日、微信、职业、地陪介绍、地陪标签等。",
                "已在页面层使用 `isAdmin`、`isGuideApproved`、`isBanned` 等能力控制功能入口，比如后台审核和地陪申请。",
                "已保留阿里云短信 Hook 的 Edge Function，说明正式短信通道已有服务端接入口。",
            ],
            "todo": [
                "正式短信能力尚未真正切到生产供应商，当前默认仍是 Supabase Phone Auth，无法视为中国大陆正式运营方案。",
                "缺实名/KYC 能力，当前用户和地陪都没有完成真正的人证校验、活体核验或公安实名校验。",
                "缺设备风险控制、频次限制、图形验证码、异常登录识别和黑名单策略。",
                "地陪身份与普通用户身份虽然已区分，但后台角色仍不完整，尚缺审核员、客服、财务、运营等更细分角色。",
            ],
            "judgement": "账号体系已经比纯原型成熟很多，但距离正式版还差“真实可管控的身份系统”。打车、租赁、到家服务类 App 真正上线时，实名、设备风控、频率限制和异常申诉几乎都是必做项。",
            "refs": [
                "lib/providers/user_provider.dart",
                "lib/models/user.dart",
                "lib/services/phone_auth_service.dart",
                "supabase/functions/send-sms-hook/index.ts",
            ],
        },
        {
            "title": "五、广场内容与招募发布",
            "summary": "广场模块已经不是单纯的静态首页，而是一条真实的内容分发链：有推荐/最新/关注三类内容视图，有帖子详情，有评论、点赞、收藏、关注流、个人发布内容和足迹。底部中间发布入口也已经分成“分享游玩瞬间”“发布招募”“发布需求”“申请成为地陪”四种动作，这一点已经比较接近社区化服务平台的实际形态。",
            "done": [
                "已完成广场主页面结构，包括城市切换、搜索、签到弹窗、推荐/最新/关注 Tab。",
                "已完成帖子 Provider，支持从 Supabase 加载帖子、搜索帖子、加载关注用户帖子、按用户查询帖子。",
                "已支持帖子点赞、帖子收藏、评论、帖子足迹记录和关注流显示。",
                "已完成帖子发布页面，并且把“普通分享”和“招募帖”做成了模式切换，而不是混成同一种内容。",
                "已完成帖子详情页和帖子卡片，说明从列表到详情的内容链已经走通。",
                "已为帖子评论单独准备 `post_comments_init.sql`，并在 Provider 中对接评论读取和写入。",
            ],
            "todo": [
                "缺内容审核机制，当前文字审核仅在地陪申请简介中有简单敏感词校验，帖子内容还缺正式审核。",
                "缺图片上传与存储治理的完整闭环，虽然有 Storage 初始化脚本，但正式图片审核、压缩、违规拦截和清理策略未形成。",
                "缺举报、拉黑、屏蔽、内容下架、热度排序、话题标签和推荐规则。",
                "招募帖虽然已经有发布入口，但围绕招募帖的列表筛选、报名、状态结束、满员关闭等规则仍未成体系。",
            ],
            "judgement": "广场已经具备演示闭环，是目前相对成熟的 C 端模块之一；但如果对标小红书式本地生活社区、租搭子、约伴平台，它还缺运营规则和内容安全能力。",
            "refs": [
                "lib/pages/home/home_page.dart",
                "lib/pages/home/post_create_page.dart",
                "lib/pages/home/post_detail_page.dart",
                "lib/providers/post_provider.dart",
                "post_comments_init.sql",
            ],
        },
        {
            "title": "六、服务页与地陪展示",
            "summary": "服务页目前承接的是“找地陪”这条链路。它已经有分类入口、城市选择、关键词搜索、地陪列表卡片、地陪详情入口、收藏/点赞/足迹数据同步，以及从地陪详情进入下单页的能力。就结构上说，它已经具备“地陪市场”的雏形。",
            "done": [
                "已完成服务页首页，包括分类图标区、城市选择、搜索框和地陪列表展示。",
                "城市选择不再只是几个写死选项，已统一接到地点选择页，可借助 A-Z 城市表和搜索结果切换城市。",
                "已完成 `GuideProvider`，包含城市过滤、关键词搜索、收藏、点赞、足迹记录和用户交互状态同步。",
                "已完成地陪卡片与地陪详情页链路，并且在下单前校验导游是否存在、是否通过审核。",
                "已在数据库侧建立 `guides / favorites / guide_likes / footprints` 等支撑表和对应 RLS。",
            ],
            "todo": [
                "缺服务价格体系和套餐体系的真实数据模型，目前更多是演示态展示。",
                "缺地陪日程、可预约时间、服务半径、接单上限、暂停接单等 B 端必备能力。",
                "缺“地陪看需求并抢单/报价/确认”的机制，当前交易更像用户直接对某个地陪下单。",
                "缺地陪信誉体系，例如完成率、超时率、取消率、投诉率、认证标签、服务保障标识。",
            ],
            "judgement": "这一模块的前台表现已经挺像一个真实 App，但内核还是展示型强、运营型弱。对标打车、租赁、到家服务平台时，这里后续一定会演化出排班、抢单、报价、服务半径和履约评分。",
            "refs": [
                "lib/pages/companion/companion_page.dart",
                "lib/pages/companion/guide_detail_page.dart",
                "lib/providers/guide_provider.dart",
                "lib/widgets/service_guide_card.dart",
            ],
        },
        {
            "title": "七、需求定制与需求池",
            "summary": "需求定制已经从原先的原型页面，发展成“可填、可发、可落库、可查看列表”的模块。需求页现在已经支持地点选择、城市联动、服务时间选择、人数与性别输入、标签选择、预算填写；发布后会写入 `demands` 表并进入需求列表页。这意味着“用户先发需求，再由平台或地陪消费需求”的产品方向已经在代码层露出雏形。",
            "done": [
                "已完成需求创建页，表单结构明显贴近原型，包括地点、时间、人数、性别、预算和标签等信息项。",
                "已完成需求列表页，具备搜索、城市筛选、推荐/最新/附近视图和发布入口。",
                "已完成 `DemandProvider`，支持从 Supabase 加载需求、创建需求、筛选需求。",
                "已补充 `demands` 表及其 RLS 策略脚本，并将发布后的去向明确到 `/demands`。",
                "已把需求地点选择接到了统一的地点选择页，支持地图、搜索和 A-Z 城市表能力的复用。",
            ],
            "todo": [
                "缺需求详情页，当前需求从列表进来后仍缺完整的详情消费链路。",
                "缺地陪报名、抢单、报价、用户确认、关闭需求、需求结束、需求取消等状态机。",
                "缺“需求和订单”的中间层，现在发需求和直接下单是两条平行线，尚未形成统一撮合逻辑。",
                "缺需求过期、置顶、推荐、审核和举报规则；平台也还没有需求运营后台。",
            ],
            "judgement": "需求模块是当前最值得继续深挖的业务点，因为它能把“找地陪”从被动浏览模式拉到主动发布模式。但它现在更像“可发布的需求板”，还不是“可运营的撮合系统”。",
            "refs": [
                "lib/pages/demand/demand_create_page.dart",
                "lib/pages/demand/demand_list_page.dart",
                "lib/providers/demand_provider.dart",
                "supabase_minimal_fix.sql",
            ],
        },
        {
            "title": "八、订单、支付、钱包与结算",
            "summary": "交易模块已经不是空白。当前有订单模型、下单页、订单列表、订单状态分类、支付宝支付服务、Supabase Edge Function 下单参数生成、支付结果回写、钱包表、交易流水表、待结算余额和完成订单后的分账逻辑。从结构上看，这一层已经从“按钮跳转”走向“交易系统搭底”。",
            "done": [
                "已完成 `Order` 模型和 `OrderStatus` 枚举，覆盖待付款、进行中、待评价、已完成、已取消等状态。",
                "已完成下单页，能从地陪详情发起订单，填服务内容、地点、时间、人数和支付方式。",
                "已完成 `OrderProvider`，支持订单加载、创建、支付、取消、完成结算和分类统计。",
                "已完成支付宝支付服务抽象，Flutter 端可通过 `PaymentService` 发起支付，并可切换沙箱/正式环境。",
                "已完成 `alipay-create-order` 与 `alipay-notify` Edge Function，前者负责生成 App 支付订单串，后者预留回调入口。",
                "已完成钱包和流水的数据库结构，以及 `increment_pending_balance`、`unfreeze_and_credit_balance` 这类 RPC。",
                "已实现取消次数累计与自动封禁的基础逻辑，说明交易侧已经开始连到风控侧。",
            ],
            "todo": [
                "支付宝回调还只是最轻量形态，当前 `alipay-notify` 只返回 `success`，还缺验签、订单核验、幂等处理和实际入账确认。",
                "缺退款、售后、改期、订单超时关闭、地陪拒单、用户取消规则、申诉仲裁等正式订单规则。",
                "缺完整的钱包前端能力，目前已有表和函数，但提现、冻结明细、资金流向、财务对账视图还不完整。",
                "缺发票、佣金规则后台配置、营销券核销、分账合规与对账报表。",
                "订单目前仍更接近 C 端单边下单，不是多方确认型服务订单。",
            ],
            "judgement": "这是当前最接近“真实商业系统”的一层，但也是最不能只看页面的一层。对标打车、租赁和到家服务平台，这里后续需要把“支付成功之后发生什么”全部明确下来，否则系统仍然只是一条可测试链路。",
            "refs": [
                "lib/models/order.dart",
                "lib/providers/order_provider.dart",
                "lib/services/payment_service.dart",
                "supabase/functions/alipay-create-order/index.ts",
                "pro_orders_init.sql",
                "pro_financial_init.sql",
            ],
        },
        {
            "title": "九、地陪申请与后台审核",
            "summary": "地陪入驻模块已经形成前后台两端：前台有申请向导，后台有待审核列表和详情处理。申请链路中已经出现实名步骤、资料填写、服务标签、合同阅读与签署、重复申请拦截、违禁词检查、封禁用户拦截、申请状态回显等特征，这已经很接近“运营后台前的最终一公里”。",
            "done": [
                "已完成多步骤地陪申请页面，包含身份、资料和合同三个阶段。",
                "已支持读取当前用户历史申请状态，已申请用户不会重复进入完整提交流程。",
                "已对封禁用户、已通过地陪、已有待审申请用户做申请拦截。",
                "已在 `ApplicationProvider` 中接入申请读写、待审列表加载和审核动作。",
                "已完成后台审核列表页和详情页入口，并在路由层对管理员做访问控制。",
                "已在 SQL 中补入 `guide_applications` 表、审核 RLS 和审核通过后同步 `guides`/创建钱包等数据库侧逻辑。",
            ],
            "todo": [
                "所谓“实名”当前还是资料提交流程的一部分，缺真实第三方实名核验、人脸识别和身份证 OCR 接口。",
                "缺审核日志、二次复审、审核备注、驳回模板、证件有效期校验和材料补交通道。",
                "后台能力还比较单点，目前主要围绕地陪审核，尚未扩展到订单、举报、风控、财务和内容管理。",
                "合同签署只是前端确认与时间写入，还不是真正的电子签约闭环。",
            ],
            "judgement": "它已经跨过了“只有按钮没有逻辑”的阶段，但如果目标是正式地陪平台，这一模块最终会演变成完整的供应侧入驻系统，既要管资质，也要管后续经营行为。",
            "refs": [
                "lib/pages/apply/apply_guide_page.dart",
                "lib/pages/admin/audit_list_page.dart",
                "lib/providers/application_provider.dart",
                "guide_applications_init.sql",
                "supabase_persistence_init.sql",
            ],
        },
        {
            "title": "十、消息、聊天室与通知",
            "summary": "消息模块当前已包含会话列表、聊天室、实时订阅、未读数统计和“获取或创建会话”的核心动作。虽然系统通知、活动通知、客服通知还没完成，但至少 IM 的主干不是空的，说明项目已经考虑到“先聊后下单”和“订单前沟通”的实际场景。",
            "done": [
                "已完成消息主页面，包括会话列表、订单服务入口、活动通知/系统通知/客服入口占位。",
                "已完成 `MessageProvider`，支持加载会话列表、进入聊天室、实时监听新消息、发送消息和未读清零。",
                "已补充 `chat_rooms`、`messages` 表及相应 RLS 策略与实时配置脚本。",
                "已支持 `getOrCreateRoom`，意味着从别的业务模块可以直接跳转进入已有或新建会话。",
            ],
            "todo": [
                "缺订单与会话的正式绑定，例如一个订单对应哪条会话、会话内是否可见订单卡片、是否可触发售后。",
                "缺系统通知、活动消息、审核结果消息、支付结果消息、客服工单消息的真实下发。",
                "缺图片消息、语音消息、已读回执、撤回、拉黑、投诉、敏感词审查等常见 IM 能力。",
                "缺客服体系和人工坐席后台，当前客服入口仍是保留位。",
            ],
            "judgement": "从代码角度看，IM 已经起步了；从产品角度看，它还只是“沟通工具”，还没成为交易和服务流程的一部分。对标打车、租赁、二手交易类 App，消息系统最终一定会与订单状态强绑定。",
            "refs": [
                "lib/pages/messages/messages_page.dart",
                "lib/pages/messages/chat_room_page.dart",
                "lib/providers/message_provider.dart",
                "pro_im_init.sql",
                "pro_im_realtime_fix.sql",
            ],
        },
        {
            "title": "十一、个人中心、资料与账户页",
            "summary": "个人中心已经不是一页静态头像和昵称，而是开始承担账户分发功能。当前已包含个人资料头部、管理员入口、客服入口、设置入口、订单入口、收藏/足迹/关注、钱包/余额/优惠券等页面分发，以及资料编辑弹窗和角色身份展示。这一层对于用户感知来说已经很像真实产品。",
            "done": [
                "已完成个人中心首页，包含用户身份展示、管理员入口、地陪身份展示和订单区块。",
                "已补充多个账户类页面：设置、安全设置、通知设置、帮助反馈、余额、钱包、优惠券、我的订单、收藏、足迹、关注。",
                "已支持资料编辑，用户资料字段明显比原始版本丰富，并可同步回 `users` 和 `guides` 表。",
                "已让订单区块可直接跳转到对应状态页，不再只能通过“更多”进入。",
            ],
            "todo": [
                "很多账户子页面仍偏演示态，例如钱包、提现、发票、客服工单、会员成长体系、评价管理等尚未完全落地。",
                "普通用户资料与地陪经营资料仍然需要进一步拆分，否则后续会在展示名、实名、经营信息之间混淆。",
                "缺实名状态、认证状态、违规记录、保证金、经营数据、接单数据等供应侧信息展示。",
            ],
            "judgement": "用户中心的外观成熟度已经不错，但要成为真正的“账户系统”，还得把钱包、评价、身份、经营和售后都收进来。",
            "refs": [
                "lib/pages/profile/profile_page.dart",
                "lib/pages/profile/orders_page.dart",
                "lib/pages/profile/settings_page.dart",
                "lib/providers/user_provider.dart",
            ],
        },
        {
            "title": "十二、地图、定位与地点选择",
            "summary": "地图能力是这次项目里很有代表性的一条底层链路。代码里已经有高德 Web Service 封装，包括地点搜索、地址地理编码、逆地理编码、当前位置获取和静态地图 URL 生成；页面侧有统一地点选择页，支持 A-Z 全国城市表、关键字搜索、当前定位和地图中心选点。对一个本地服务 App 来说，这是非常关键的一层，因为它影响广场城市、服务城市、订单地点和需求地点的统一性。",
            "done": [
                "已完成 `MapService` 抽象与 `AmapMapService` 实现，包含搜索、正逆地理、当前位置和静态图能力。",
                "已完成统一地点选择页，能够被订单、需求、广场和服务页复用。",
                "已内置全国城市 A-Z 列表，并支持城市搜索，不再局限于少量预设城市。",
                "已将广场、服务页和需求页的城市/地点选择逐步统一到同一个逻辑源头。",
            ],
            "todo": [
                "当前主要依赖高德 Web Service，若要做更真实的移动端定位体验，Android/iOS 原生定位与地图 SDK 仍需继续对齐。",
                "手机端地图显示和定位稳定性仍需继续真机联调，尤其是权限、网络、静态图和真实瓦片地图的兼容性。",
                "缺地点收藏、常用地址、服务半径、附近推荐和 POI 落点确认等更深一层的地图产品能力。",
            ],
            "judgement": "地图底座已经搭得很不错，但正式版会进一步分成“高德 Web Service 负责查询”和“原生/SDK 负责定位与地图交互”两层，这样稳定性和体验都会更好。",
            "refs": [
                "lib/services/map_service.dart",
                "lib/pages/order/location_picker_page.dart",
                "lib/config/amap_config.dart",
            ],
        },
        {
            "title": "十三、数据库、RLS、Edge Functions 与存储",
            "summary": "这部分是当前仓库最容易被低估、但其实最关键的成果之一。根目录已经有大量 SQL 初始化和补丁脚本，覆盖基础用户关系、帖子、地陪、需求、订单、钱包、IM、收藏、足迹、评论、存储权限等；同时还有 3 个 Supabase Edge Function，分别针对支付宝下单、支付宝回调和短信发送。这意味着项目已经不只是 Flutter 前端，而是开始具备一套 BaaS 驱动的小型后端系统。",
            "done": [
                "已建立基础业务表脚本，包括 `users`、`posts`、`guides`、`favorites`、`footprints`、`guide_likes`、`follows` 等。",
                "已补地陪申请、订单、需求、评论、钱包、交易流水、聊天会话和消息表。",
                "已为多个表补上 RLS 策略，至少开始考虑“谁能看、谁能写”的真实权限边界。",
                "已补多类 RPC / 触发器，例如取消次数累计、自动封禁、审核通过同步地陪、创建钱包、评论计数、待结算余额处理。",
                "已建立 Supabase Storage 的基础策略，为帖子图片和头像上传做了权限准备。",
                "已编写 3 个 Edge Function：`alipay-create-order`、`alipay-notify`、`send-sms-hook`。",
            ],
            "todo": [
                "SQL 脚本目前存在“初始化脚本 + 补丁脚本并存”的状态，后续需要做一次版本整理，避免不同环境执行顺序不一致。",
                "部分业务规则仍主要靠前端守，服务端仍需加强幂等、状态校验和权限兜底。",
                "支付回调、短信回调、存储清理、日志审计、风控事件记录等服务端能力仍然较轻。",
                "如果后续迁移到自建服务器，需要明确哪些能力继续留在 Supabase，哪些迁出到独立后端。",
            ],
            "judgement": "后端底座已经有了，但还没有真正完成“生产后端治理”。从试运营走向正式运营时，这一层会是风险最高也最值得继续投入的一层。",
            "refs": [
                "supabase_init.sql",
                "supabase_persistence_init.sql",
                "supabase_minimal_fix.sql",
                "pro_orders_init.sql",
                "pro_financial_init.sql",
                "pro_im_init.sql",
                "supabase/functions/alipay-create-order/index.ts",
                "supabase/functions/send-sms-hook/index.ts",
            ],
        },
        {
            "title": "十四、运维部署与客户交付文档",
            "summary": "项目根目录已经出现了一批明显面向交付和实施的文档，例如 `PROJECT_DEPLOYMENT_AND_API_SETUP_GUIDE.md`、`CLIENT_HANDOFF_STEPS.md`、`CLIENT_DELIVERABLES.md`、`EXTERNAL_SETUP_LINKS.md`。这说明项目已经从“自己本地调通”进入“准备和客户、第三方平台、服务器环境协作”的阶段。",
            "done": [
                "已整理客户需要提供的主体、地图 Key、支付 Key、短信资质、服务器和备案材料。",
                "已整理阿里云 ECS / RDS / 域名 / HTTPS / ICP / App 备案的详细操作步骤。",
                "已明确正式部署建议结构：Flutter App -> Nginx -> 自建后端 API -> PostgreSQL。",
                "已将第三方接入步骤文档化，便于后续与客户直接对接。",
            ],
            "todo": [
                "文档虽全，但正式环境尚未真正搭起，服务器、域名、备案、证书和 CI/CD 仍需落地。",
                "缺统一的部署脚本、环境变量模板、数据库迁移顺序说明和发布回滚策略。",
                "缺监控、日志、告警、备份恢复、对象存储生命周期等正式运维能力。",
            ],
            "judgement": "这说明项目推进方式已经开始工程化，但部署系统本身还没有形成。也就是说，文档准备得比生产环境更靠前，这是好事，但后面要尽快把文档转成真实环境。",
            "refs": [
                "PROJECT_DEPLOYMENT_AND_API_SETUP_GUIDE.md",
                "CLIENT_HANDOFF_STEPS.md",
                "CLIENT_DELIVERABLES.md",
                "EXTERNAL_SETUP_LINKS.md",
            ],
        },
    ]

    process_sections = [
        (
            "十五、对标类似 App 后仍然必须补齐的完整流程",
            "这一部分不局限于当前仓库里已经做过什么，而是按类似打车、租赁、本地生活服务、导游预约平台的完整业务路径，反推一个地陪 App 真正要跑起来还需要哪些环节。它的目的不是增加工作量，而是防止后续只补页面、不补流程。",
            [],
        ),
        (
            "十六、C 端用户完整链路仍需补齐的内容",
            "如果把普通用户看作消费者，完整链路至少要覆盖“进入平台 -> 信任平台 -> 找人/发需求 -> 成交 -> 履约 -> 售后 -> 复购”。当前项目前半段已经形成，但中后段还缺很多正式规则。",
            [
                "注册与登录：补正式短信验证码、异常设备拦截、注册频控、图形验证码、实名状态展示。",
                "浏览与发现：补标签/话题/附近推荐/热榜、举报拉黑、内容审核、搜索排序。",
                "发需求与下单：补需求详情、报名/抢单/报价、订单确认页、优惠券和规则提示。",
                "支付与履约：补支付成功确认、服务开始确认、到场确认、超时未到、改期、取消费规则。",
                "售后与沉淀：补评价、投诉、退款、申诉、客服工单、订单追踪、常用地点和历史服务记录。",
            ],
        ),
        (
            "十七、地陪端完整链路仍需补齐的内容",
            "如果把地陪看作供应侧，当前项目只完成了“展示 + 申请 + 被下单”的一部分。参考打车司机端、租赁平台商家端和到家服务技师端，地陪端后续至少需要下面这些能力。",
            [
                "入驻前：实名、证件、合同、服务协议、头像与简介审核、标签和擅长领域管理。",
                "接单前：可接单状态、服务城市、服务半径、可预约时间、价格体系、排班日历。",
                "接单中：需求报名、报价、确认、拒单、改期、超时未响应、系统派单/抢单。",
                "履约中：到场确认、服务开始、临时沟通、异常上报、结束确认。",
                "经营中：钱包、待结算、提现、月度收入、评分、投诉率、接单率、取消率、成长体系。",
            ],
        ),
        (
            "十八、平台管理端完整链路仍需补齐的内容",
            "一个地陪 App 不可能只靠前台页面运营。参考打车平台、民宿平台、租赁平台，它至少要有后台来处理内容、订单、风控、客服和资金。",
            [
                "账号与角色后台：管理员、审核员、客服、财务、运营的角色拆分与权限矩阵。",
                "内容后台：帖子审核、招募帖审核、需求审核、举报处理、敏感词和黑名单配置。",
                "供应侧后台：地陪申请审核、地陪封禁、资质复核、服务质量统计、违规处理。",
                "交易后台：订单状态总览、异常订单、支付对账、退款审核、售后工单、佣金配置。",
                "运营后台：轮播位、推荐位、城市活动、消息推送、优惠券、营销活动。",
                "客服后台：会话分配、投诉记录、处理结果、回访、申诉闭环。",
            ],
        ),
        (
            "十九、资金、风控与合规链路仍需补齐的内容",
            "这一部分通常是产品从“能用”到“能上线”的分水岭。越像真实交易平台，这一层越不能省。",
            [
                "支付：正式支付宝回调验签、幂等入账、退款、对账、支付异常补偿。",
                "资金：钱包明细、提现申请、提现审核、平台佣金报表、发票与税务处理。",
                "风控：账号风险、设备风险、异常取消、私下交易识别、敏感词、恶意投诉识别。",
                "合规：实名认证、隐私协议、服务协议、地陪责任说明、App 备案、服务器备案。",
                "日志审计：后台操作日志、订单状态变更日志、支付日志、审核日志、客服日志。",
            ],
        ),
    ]

    priorities = [
        ("P0", "把交易主链做成闭环", "明确需求撮合路径；梳理订单状态机；补支付回调验签；补取消/退款/售后规则；让订单、消息、履约互相可追踪。", "支付宝正式能力、服务器环境、域名 HTTPS、客户主体"),
        ("P1", "把供应侧做成可运营", "补地陪实名认证、资质审核、排班、接单/抢单、报价、可接单状态、经营数据。", "实名/审核相关 API、客户审核规则、后台角色设计"),
        ("P2", "把后台和风控搭起来", "补内容审核、举报处理、订单后台、客服工单、财务视图、日志审计、异常账号处理。", "管理端需求确认、运维资源、数据库脚本整理"),
        ("P3", "把基础设施从测试态迁到生产态", "完成服务器、数据库、对象存储、短信正式通道、CI/CD、监控告警、备份恢复。", "阿里云服务器、RDS、短信资质、域名、备案、SSL"),
    ]

    client_items = [
        "项目主体：明确最终以上线企业、个体户还是个人主体作为支付宝、高德、短信和备案的归属方。",
        "地图能力：提供高德 Web Service Key；如需真地图定位体验，继续提供 Android / iOS 地图 SDK 所需配置。",
        "支付能力：提供支付宝 AppID、应用公钥配置结果、支付宝公钥、正式回调域名与商户主体信息。",
        "短信能力：提供正式短信供应商账号、签名、模板、AccessKey/SecretKey 和主体资质。",
        "服务器与网络：提供 ECS、RDS、域名、HTTPS 证书、备案信息、测试环境与正式环境地址。",
        "合规材料：提供用户协议、隐私政策、客服联系方式、投诉处理规则、平台服务协议。",
        "运营规则：拍板订单取消规则、退款规则、地陪抽佣规则、审核规则、投诉处置规则。",
    ]

    document = Document()
    for section in document.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)
        section.header_distance = Cm(1.2)
        section.footer_distance = Cm(1.2)

    build_styles(document)
    add_cover(document)
    add_overview_metrics(document, metrics)
    add_status_table(document, status_rows)

    for module in modules:
        document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
        add_module_section(document, module)

    for title, intro, items in process_sections:
        document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
        if items:
            add_process_section(document, title, intro, items)
        else:
            document.add_heading(title, level=1)
            add_para(document, intro)

    add_priority_table(document, priorities)
    document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    add_client_section(document, client_items)
    add_conclusion(document)

    # Footer
    section = document.sections[0]
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer.add_run("伴一下 App 项目整理报告 | 自动生成于 " + str(date.today()))
    set_run_font(footer_run, size=9, color="5B657A")

    document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
